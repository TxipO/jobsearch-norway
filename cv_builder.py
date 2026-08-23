"""Deterministic CV + søknad builder. Python owns the entire document
structure and all the facts; the only thing a per-vacancy tailoring step
can vary is the small set of text fields in `tailoring` (summary, which
jobs are Relevant vs Other, tool ordering, søknad paragraphs). This is the
core fix from PLAN-BUILDER.md: the model physically cannot break the
Norwegian CV structure or editorialize in a section it doesn't control,
because it never writes the document — it returns data, this builds it.

Structure follows a real successful Norwegian CV (see PLAN-BUILDER.md
finding #2): two-column contact header, summary, Relevant erfaring / Annen
erfaring split, Utdanning with Norwegian degree equivalents, Programverktøy,
Språk, and an optional Hobbyer line. English throughout (candidate's honest
strongest language).
"""

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

import profile_data as pd

REPO_ROOT = Path(__file__).parent
PROFILE_DIR = REPO_ROOT / "profile"
PERSONAL_PATH = PROFILE_DIR / "personal.json"

MUTED = RGBColor(0x55, 0x55, 0x55)
RIGHT_TAB = Inches(6.5)  # matches the usable width inside 0.9" side margins on US Letter

# Section heading translations. Norwegian terms chosen to match the real
# successful CV this builder's structure is based on (PLAN-BUILDER.md) and
# cv-reference.md's own vocabulary ("personlig egnethet" -> Personlige
# egenskaper for the soft-skills section).
SECTION_LABELS = {
    "en": {
        "profile": "Profile",
        "relevant": "Relevant experience", "other": "Other experience", "experience": "Experience",
        "education": "Education", "hard_skills": "Hard Skills", "soft_skills": "Soft Skills",
        "languages": "Languages", "interests": "Interests",
    },
    "no": {
        "profile": "Profil",
        "relevant": "Relevant erfaring", "other": "Annen erfaring", "experience": "Erfaring",
        "education": "Utdanning", "hard_skills": "Faglige ferdigheter", "soft_skills": "Personlige egenskaper",
        "languages": "Språk", "interests": "Interesser",
    },
}


def load_personal() -> dict:
    if PERSONAL_PATH.exists():
        return json.loads(PERSONAL_PATH.read_text(encoding="utf-8"))
    # Fall back to placeholders so the builder still runs before personal.json
    # is filled in — the output just carries visible [brackets] to replace.
    return {
        "name": "[Your Name]", "phone": "[phone]", "email": "[email]",
        "address_line": "[Address], Norway", "linkedin": "", "hobbies": [],
    }


def _add_bottom_border(p, color="999999", size="4") -> None:
    """Thin horizontal rule under a paragraph — shared by _section_heading
    and the søknad header divider, so both documents use the exact same
    visual device instead of two hand-rolled approximations of it."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def _apply_letter_spacing(run, points: float) -> None:
    """OOXML character spacing (`w:spacing`, in twentieths of a point) has
    no public python-docx API — set the raw XML directly. The reference
    resumes' uppercase labels (section headings, role headline) all use
    CSS letter-spacing; plain bold-caps text without it looks visibly
    tighter/more generic side by side (2026-07-20, user feedback that the
    generated CVs looked less polished than their own)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rPr = run._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), str(round(points * 20)))
    rPr.append(spacing)


def _section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(3)
    # Without this, a heading can land as the last line on a page with its
    # content pushed to the next — caught live in the rendered PDF,
    # 2026-07-17 (Hard Skills heading orphaned from its own list).
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    _apply_letter_spacing(run, 0.6)
    _add_bottom_border(p)
    return p


def _localize_dates(dates: str, lang: str) -> str:
    """"dates" is otherwise treated as language-agnostic (numbers/month
    abbreviations don't need translating), but "present" is an actual
    English word, not a proper noun — found live 2026-07-19 still showing
    "Nov 2024 – present" in the NO-language CV (both the docx pipeline and
    a one-off HTML export), never caught because nothing exercised the
    output text end-to-end for this specific word before."""
    if lang == "no":
        return re.sub(r"\bpresent\b", "nå", dates, flags=re.IGNORECASE)
    return dates


def _job_entry(doc, job, lang="en"):
    """Underlined 'Title, Company' on the left, dates right-aligned on the
    same line via a right tab stop, then the description as flowing prose
    (not bullets) — the real-CV pattern, not our old bullet lists."""
    title = job["title_no"] if lang == "no" else job["title"]
    location = job.get("location_no") if lang == "no" else job.get("location")
    description = job["description_no"] if lang == "no" else job["description"]

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
    head = p.add_run(f"{title}, {job['company']}")
    head.bold = True
    head.underline = True
    p.add_run("\t")
    dates = p.add_run(_localize_dates(job["dates"], lang))
    dates.italic = True
    dates.font.color.rgb = MUTED

    if location:
        loc = doc.add_paragraph()
        loc.paragraph_format.space_after = Pt(1)
        lr = loc.add_run(location)
        lr.font.size = Pt(9)
        lr.font.color.rgb = MUTED

    body = doc.add_paragraph(description)
    body.paragraph_format.space_after = Pt(4)


def _body_line(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    return p


def _set_default_font(doc: Document, name: str) -> None:
    """style.font.name alone only sets w:rFonts/@w:ascii — LibreOffice (and
    sometimes Word) render Latin text with diacritics using @w:hAnsi
    instead, which is left at its theme default if unset. Verified live
    2026-07-19: .font.name = "Segoe UI" alone rendered as serif in the
    LibreOffice-exported PDF (æøå-heavy Norwegian text apparently doesn't
    count as pure-ASCII for this purpose) — setting ascii/hAnsi/cs
    explicitly via the raw XML is the actual fix."""
    from docx.oxml.ns import qn
    style = doc.styles["Normal"]
    style.font.name = name
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), name)


def _new_cv_document() -> Document:
    doc = Document()
    # The reference documents (Claude Design PDF exports) use Libre
    # Franklin, a Google Font unlikely to be installed on whatever machine
    # opens this .docx — Word/LibreOffice would silently substitute an
    # unpredictable fallback rather than error, which defeats the point.
    # Segoe UI ships with every modern Windows install and is the closest
    # widely-available match to Libre Franklin's clean geometric-humanist
    # character (2026-07-19, user-requested visual match to those PDFs).
    _set_default_font(doc, "Segoe UI")
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
    return doc


def _resolve_photo_path(personal: dict) -> Path | None:
    photo_path = personal.get("photo_path")
    if not photo_path:
        return None
    candidate = Path(photo_path)
    if not candidate.is_absolute():
        candidate = REPO_ROOT / candidate
    return candidate if candidate.exists() else None


def _write_header_plain(doc, personal: dict, role_headline: str) -> None:
    name_p = doc.add_paragraph()
    name_p.paragraph_format.space_after = Pt(2)
    nr = name_p.add_run(personal["name"])
    nr.bold = True
    nr.font.size = Pt(18)

    if role_headline:
        rp = doc.add_paragraph()
        rp.paragraph_format.space_after = Pt(4)
        rr = rp.add_run(role_headline.upper())
        rr.bold = True
        rr.font.size = Pt(10)
        rr.font.color.rgb = MUTED
        _apply_letter_spacing(rr, 0.8)

    # Contact: left details, right-aligned address on the same line
    contact_bits = [personal.get("email"), personal.get("phone"), personal.get("linkedin")]
    contact_line = "  ·  ".join(b for b in contact_bits if b)
    cp = doc.add_paragraph()
    cp.paragraph_format.space_after = Pt(0)
    cp.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
    cl = cp.add_run(contact_line)
    cl.font.size = Pt(10)
    cp.add_run("\t")
    ar = cp.add_run(personal.get("address_line", ""))
    ar.font.size = Pt(10)


def _write_header_with_photo(doc, personal: dict, role_headline: str, photo_path: Path) -> None:
    """Photo | name+role | contact, side by side — matches the user's own
    reference resumes (2026-07-20). Word tables are the only way python-docx
    can lay out text side-by-side, which means this content lives outside
    doc.paragraphs (table cell paragraphs aren't in that collection) — fine
    here since nothing asserts on header text, but worth knowing if a future
    test needs to reach it."""
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [Inches(1.0), Inches(3.3), Inches(2.4)]
    for col, width in zip(table.columns, widths):
        col.width = width
    photo_cell, name_cell, contact_cell = table.rows[0].cells
    for cell, width in zip((photo_cell, name_cell, contact_cell), widths):
        cell.width = width
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    photo_run = photo_cell.paragraphs[0].add_run()
    photo_run.add_picture(str(photo_path), width=Inches(0.9))

    name_p = name_cell.paragraphs[0]
    name_p.paragraph_format.space_after = Pt(2)
    nr = name_p.add_run(personal["name"])
    nr.bold = True
    nr.font.size = Pt(18)

    if role_headline:
        role_p = name_cell.add_paragraph()
        rr = role_p.add_run(role_headline.upper())
        rr.bold = True
        rr.font.size = Pt(10)
        rr.font.color.rgb = MUTED
        _apply_letter_spacing(rr, 0.8)

    contact_lines = [personal.get("email"), personal.get("phone"), personal.get("address_line")]
    first = True
    for line in contact_lines:
        if not line:
            continue
        p = contact_cell.paragraphs[0] if first else contact_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(line)
        r.font.size = Pt(10)
        first = False


def _write_header(doc, personal: dict, summary: str, role_headline: str = "", lang: str = "en") -> None:
    photo_path = _resolve_photo_path(personal)
    if photo_path:
        _write_header_with_photo(doc, personal, role_headline, photo_path)
    else:
        _write_header_plain(doc, personal, role_headline)

    # Divider between the header block and the body — same thin rule
    # _section_heading uses, matching the reference resumes' header/body
    # separation (2026-07-20).
    divider = doc.add_paragraph()
    divider.paragraph_format.space_before = Pt(2)
    divider.paragraph_format.space_after = Pt(1)
    _add_bottom_border(divider)

    # The reference resumes label the summary paragraph "Profile" — the
    # generated CV was missing this heading entirely (summary just floated
    # under the divider with no label), a real structural gap spotted
    # 2026-07-20 comparing side by side, not a styling nitpick.
    _section_heading(doc, SECTION_LABELS[lang]["profile"])
    _body_line(doc, summary).paragraph_format.space_before = Pt(0)


def _write_footer_sections(doc, hard_skills: list[str], soft_skills: list[str], personal: dict, lang="en") -> None:
    """Education / Hard Skills / Soft Skills / Languages / Interests —
    identical across every CV variant, doesn't depend on the relevant/other
    split or job ordering."""
    labels = SECTION_LABELS[lang]

    _section_heading(doc, labels["education"])
    for edu in pd.EDUCATION:
        degree = edu["degree_no"] if lang == "no" else edu["degree"]
        institution = (edu.get("institution_no") if lang == "no" else None) or edu["institution"]
        note = (edu.get("note_no") if lang == "no" else edu.get("note")) or ""

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
        dr = p.add_run(degree)
        dr.bold = True
        p.add_run("\t")
        yr = p.add_run(edu["dates"])
        yr.italic = True
        yr.font.color.rgb = MUTED
        inst = doc.add_paragraph()
        inst.paragraph_format.space_after = Pt(0)
        ir = inst.add_run(institution)
        ir.font.size = Pt(10)
        if note:
            nt = doc.add_paragraph()
            nt.paragraph_format.space_after = Pt(4)
            ntr = nt.add_run(note)
            ntr.font.size = Pt(9)
            ntr.font.color.rgb = MUTED

    _section_heading(doc, labels["hard_skills"])
    _body_line(doc, "  ·  ".join(hard_skills))

    _section_heading(doc, labels["soft_skills"])
    _body_line(doc, "  ·  ".join(soft_skills))

    _section_heading(doc, labels["languages"])
    languages = pd.LANGUAGES_NO if lang == "no" else pd.LANGUAGES
    _body_line(doc, "   ·   ".join(f"{name} — {level}" for name, level in languages))

    hobbies = (personal.get("hobbies_no") if lang == "no" else None) or personal.get("hobbies") or []
    if hobbies:
        _section_heading(doc, labels["interests"])
        _body_line(doc, ", ".join(hobbies))


def build_cv(tailoring: dict, out_path: Path, lang: str = "en") -> Path:
    """The IT-support-focused CV: Relevant/Other experience split. Aimed at
    technically-inclined roles — see build_cv_general() for the broad,
    single-timeline variant used for retail/warehouse/non-technical
    applications. lang="no" renders the Norwegian variant (cv-reference.md
    recommends having both, sent depending on the posting's language)."""
    personal = load_personal()
    summary = tailoring.get("summary") or (pd.DEFAULT_SUMMARY_NO if lang == "no" else pd.DEFAULT_SUMMARY)
    relevant = tailoring.get("relevant_job_ids") or pd.DEFAULT_RELEVANT
    other = tailoring.get("other_job_ids") or pd.DEFAULT_OTHER
    hard_skills = tailoring.get("tools_highlight") or (pd.HARD_SKILLS_NO if lang == "no" else pd.HARD_SKILLS)
    soft_skills = tailoring.get("soft_skills_highlight") or (pd.SOFT_SKILLS_NO if lang == "no" else pd.SOFT_SKILLS)
    role_headline = tailoring.get("role_headline") or (pd.ROLE_HEADLINE_NO if lang == "no" else pd.ROLE_HEADLINE)
    labels = SECTION_LABELS[lang]

    doc = _new_cv_document()
    _write_header(doc, personal, summary, role_headline, lang)

    _section_heading(doc, labels["relevant"])
    for jid in relevant:
        if jid in pd.JOBS:
            _job_entry(doc, pd.JOBS[jid], lang)

    if other:
        _section_heading(doc, labels["other"])
        for jid in other:
            if jid in pd.JOBS:
                _job_entry(doc, pd.JOBS[jid], lang)

    _write_footer_sections(doc, hard_skills, soft_skills, personal, lang)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def build_cv_general(tailoring: dict, out_path: Path, lang: str = "en") -> Path:
    """Broad-applications CV: single reverse-chronological "Experience"
    list, no relevant/other split — for retail, warehouse, and other
    non-technical applications where the IT-support framing doesn't fit
    and splitting the timeline would look like padding, not focus."""
    personal = load_personal()
    summary = tailoring.get("summary") or (pd.GENERAL_SUMMARY_NO if lang == "no" else pd.GENERAL_SUMMARY)
    hard_skills = tailoring.get("tools_highlight") or (pd.HARD_SKILLS_NO if lang == "no" else pd.HARD_SKILLS)
    soft_skills = tailoring.get("soft_skills_highlight") or (pd.SOFT_SKILLS_NO if lang == "no" else pd.SOFT_SKILLS)
    role_headline = tailoring.get("role_headline") or (pd.GENERAL_ROLE_HEADLINE_NO if lang == "no" else pd.GENERAL_ROLE_HEADLINE)
    labels = SECTION_LABELS[lang]

    doc = _new_cv_document()
    _write_header(doc, personal, summary, role_headline, lang)

    _section_heading(doc, labels["experience"])
    for jid in pd.ALL_JOBS_CHRONOLOGICAL:
        if jid in pd.JOBS:
            _job_entry(doc, pd.JOBS[jid], lang)

    _write_footer_sections(doc, hard_skills, soft_skills, personal, lang)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def build_soknad(tailoring: dict, out_path: Path, lang: str = "en") -> Path:
    """Same visual grammar as build_cv (see _new_cv_document/_write_header/
    _section_heading) — user-requested 2026-07-19: the søknad used to be
    Word's plain default styling (Calibri, no rule, its own one-off
    margins) while the CV had a deliberate typographic system. A søknad and
    its paired CV are read back to back by the same employer; they should
    visibly belong to the same document set, not look like two different
    tools made them.

    lang="no" (2026-07-20): per cv-reference.md's own research on Norwegian
    søknad convention, there's no "Dear Hiring Team,"-equivalent opener —
    Norwegian formal letters go straight from the position-line heading into
    the body. Closes with "Med vennlig hilsen," (cv-reference.md section 4),
    not a translated "Kind regards,"."""
    personal = load_personal()
    soknad = tailoring.get("soknad") or {}
    default_position = "Søknad" if lang == "no" else "Application"
    position_line = soknad.get("position_line", default_position)
    paragraphs = soknad.get("paragraphs") or []

    doc = _new_cv_document()

    name_p = doc.add_paragraph()
    name_p.paragraph_format.space_after = Pt(2)
    nr = name_p.add_run(personal["name"])
    nr.bold = True
    nr.font.size = Pt(18)

    position_p = doc.add_paragraph()
    position_p.paragraph_format.space_after = Pt(6)
    pr = position_p.add_run(position_line)
    pr.font.size = Pt(11)
    pr.font.color.rgb = MUTED

    contact_bits = [personal.get("email"), personal.get("phone"), personal.get("linkedin")]
    contact_line = "  ·  ".join(b for b in contact_bits if b)
    cp = doc.add_paragraph()
    cp.paragraph_format.space_after = Pt(0)
    cp.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
    cl = cp.add_run(contact_line)
    cl.font.size = Pt(10)
    cl.font.color.rgb = MUTED
    cp.add_run("\t")
    ar = cp.add_run(personal.get("address_line", ""))
    ar.font.size = Pt(10)
    ar.font.color.rgb = MUTED
    # Same thin rule _section_heading uses under every CV section — the
    # divider between the header block and the letter body, matching how
    # the reference documents separate header from content.
    cp.paragraph_format.space_after = Pt(10)
    _add_bottom_border(cp)

    if lang == "no":
        first = True
        for para in paragraphs:
            p = _body_line(doc, para)
            if first:
                p.paragraph_format.space_before = Pt(4)
                first = False
            p.paragraph_format.space_after = Pt(8)
        _body_line(doc, "Med vennlig hilsen,").paragraph_format.space_before = Pt(8)
    else:
        _body_line(doc, "Dear Hiring Team,").paragraph_format.space_before = Pt(4)
        for para in paragraphs:
            _body_line(doc, para).paragraph_format.space_after = Pt(8)
        _body_line(doc, "Kind regards,").paragraph_format.space_before = Pt(8)
    doc.add_paragraph(personal["name"])

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
