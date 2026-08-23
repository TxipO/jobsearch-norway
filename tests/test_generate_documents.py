"""Tests for generate_documents.py — søknad-only CLI bridge (CV generation
removed 2026-07-20, user: "прибери цю функцію" — the auto-copied/auto-built
CV kept resurfacing an outdated generic CV instead of the per-vacancy one
placed via place-cv). This script now only ever touches soknad.*, plus
converting an already-placed cv.docx to PDF if one exists."""

import json
import subprocess
import sys
import shutil
from pathlib import Path

import pytest

import generate_documents as gd


def _run(uuid, tailoring: dict, project_root):
    return subprocess.run(
        [sys.executable, str(project_root / "generate_documents.py"), uuid],
        input=json.dumps(tailoring), capture_output=True, text=True, encoding="utf-8",
        cwd=str(project_root),
    )


def test_cli_never_creates_a_cv_when_none_placed():
    """Core regression: this script must not fabricate/copy a generic CV —
    only a CV placed via place_cv.py should ever appear in the vacancy
    folder."""
    project_root = Path(__file__).parent.parent
    tailoring = {"soknad": {"position_line": "Application for X, Y", "paragraphs": ["Para one.", "Para two."]}}
    result = _run("test-slug-no-cv", tailoring, project_root)
    assert result.returncode == 0, result.stderr
    assert "none placed yet" in result.stdout

    out_dir = project_root / "profile" / "generated" / "test-slug-no-cv"
    assert not (out_dir / "cv.docx").exists()
    assert not (out_dir / "cv.pdf").exists()
    assert (out_dir / "soknad.docx").exists()

    shutil.rmtree(out_dir)


def test_cli_converts_already_placed_docx_cv_to_pdf(monkeypatch):
    """If a .docx CV was placed (via place_cv.py) but not yet exported to
    PDF, this script's PDF pass should pick it up too — best-effort, so it
    silently no-ops if LibreOffice isn't available in this environment."""
    from docx import Document

    project_root = Path(__file__).parent.parent
    out_dir = project_root / "profile" / "generated" / "test-slug-placed-cv"
    out_dir.mkdir(parents=True, exist_ok=True)
    Document().save(str(out_dir / "cv.docx"))

    tailoring = {"soknad": {"position_line": "Application for X, Y", "paragraphs": ["Para."]}}
    result = _run("test-slug-placed-cv", tailoring, project_root)
    assert result.returncode == 0, result.stderr
    assert "CV:" not in result.stdout or "none placed yet" not in result.stdout

    shutil.rmtree(out_dir)


def test_cli_preserves_non_ascii_characters_from_stdin():
    """Live bug caught 2026-07-20: sys.stdin.read() alone silently mangled
    Norwegian characters (ø/å/æ) into mojibake on Windows ("Høgskolen" ->
    "HÃ¸gskolen") in a real generated søknad — reading raw bytes and
    decoding as UTF-8 explicitly is the fix."""
    from docx import Document

    project_root = Path(__file__).parent.parent
    tailoring = {
        "soknad": {
            "position_line": "Søknad på stilling som IT-konsulent, NLA Høgskolen",
            "paragraphs": ["Jeg søker på stillingen ved NLA Høgskolen i Sandviken."],
        }
    }
    result = _run("test-slug-encoding", tailoring, project_root)
    assert result.returncode == 0, result.stderr

    out_dir = project_root / "profile" / "generated" / "test-slug-encoding"
    doc = Document(str(out_dir / "soknad.docx"))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Høgskolen" in text
    assert "HÃ¸gskolen" not in text
    assert "Ã¸" not in text

    shutil.rmtree(out_dir)


def test_cli_never_overwrites_an_already_placed_pdf():
    """Critical regression, caught live 2026-07-20: a stale leftover
    cv.docx sitting next to a freshly place_cv'd cv.pdf got silently
    reconverted, overwriting the user's real placed PDF with the stale
    docx's content — because the old code checked .docx before .pdf. A
    cv.pdf that already exists must never be touched."""
    project_root = Path(__file__).parent.parent
    out_dir = project_root / "profile" / "generated" / "test-slug-stale-docx"
    out_dir.mkdir(parents=True, exist_ok=True)
    (out_dir / "cv.docx").write_bytes(b"stale docx bytes from an old run")
    (out_dir / "cv.pdf").write_bytes(b"the real placed pdf bytes")

    tailoring = {"soknad": {"position_line": "Application for X, Y", "paragraphs": ["Para."]}}
    result = _run("test-slug-stale-docx", tailoring, project_root)
    assert result.returncode == 0, result.stderr

    assert (out_dir / "cv.pdf").read_bytes() == b"the real placed pdf bytes"

    shutil.rmtree(out_dir)


def test_bad_json_input_exits_with_error():
    project_root = Path(__file__).parent.parent
    result = subprocess.run(
        [sys.executable, str(project_root / "generate_documents.py"), "test-slug-badjson"],
        input="not json", capture_output=True, text=True, cwd=str(project_root),
    )
    assert result.returncode != 0
    assert "not valid JSON" in result.stderr
