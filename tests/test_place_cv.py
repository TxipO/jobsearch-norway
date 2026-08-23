"""Tests for place_cv.py — files a per-vacancy CV the user hand-crafted
elsewhere (2026-07-20, user-requested workflow change: CVs will now vary
per vacancy instead of one rarely-changing master)."""

import sys
from pathlib import Path

import pytest
from docx import Document

import place_cv


def _fake_vacancy_present(monkeypatch, present=True):
    monkeypatch.setattr(place_cv.db, "connect", lambda: object())
    monkeypatch.setattr(place_cv.db, "get_vacancy", lambda conn, uuid: ({"uuid": uuid} if present else None))


def test_unknown_uuid_raises(tmp_path, monkeypatch):
    _fake_vacancy_present(monkeypatch, present=False)
    src = tmp_path / "cv.docx"
    Document().save(str(src))
    with pytest.raises(ValueError, match="no vacancy"):
        place_cv.place_cv("does-not-exist", src)


def test_unsupported_extension_raises(tmp_path, monkeypatch):
    _fake_vacancy_present(monkeypatch)
    src = tmp_path / "cv.txt"
    src.write_text("hello")
    with pytest.raises(ValueError, match="unsupported"):
        place_cv.place_cv("some-uuid", src)


def test_places_docx_and_extracts_text(tmp_path, monkeypatch):
    _fake_vacancy_present(monkeypatch)
    monkeypatch.setattr(place_cv, "OUT_ROOT", tmp_path / "generated")

    doc = Document()
    doc.add_paragraph("A distinctive sentence only in this dropped-in CV.")
    src = tmp_path / "source-cv.docx"
    doc.save(str(src))

    dest = place_cv.place_cv("vac-123", src)

    assert dest == tmp_path / "generated" / "vac-123" / "cv.docx"
    assert dest.exists()
    cv_txt = dest.parent / "cv.txt"
    assert cv_txt.exists()
    assert "A distinctive sentence only in this dropped-in CV." in cv_txt.read_text(encoding="utf-8")


def test_places_pdf_as_is_and_extracts_text(tmp_path, monkeypatch):
    """PDF stays PDF — no forced docx conversion, per the user's explicit
    ask (2026-07-20) to keep the artifact they made as-is."""
    _fake_vacancy_present(monkeypatch)
    monkeypatch.setattr(place_cv, "OUT_ROOT", tmp_path / "generated")

    from pypdf import PdfWriter
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    src = tmp_path / "source-cv.pdf"
    with open(src, "wb") as f:
        writer.write(f)

    dest = place_cv.place_cv("vac-456", src)

    assert dest == tmp_path / "generated" / "vac-456" / "cv.pdf"
    assert dest.exists()
    assert (dest.parent / "cv.txt").exists()


def test_main_warns_when_extraction_is_near_empty(tmp_path, monkeypatch, capsys):
    """Live bug caught 2026-07-20: a PDF with no real selectable text layer
    (e.g. vector-outlined fonts) extracted to almost nothing via pypdf, and
    the user only found out indirectly, much later, from a badly-grounded
    søknad. main() should warn immediately at placement time instead."""
    _fake_vacancy_present(monkeypatch)
    monkeypatch.setattr(place_cv, "OUT_ROOT", tmp_path / "generated")

    from pypdf import PdfWriter
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)  # no text content at all
    src = tmp_path / "source-cv.pdf"
    with open(src, "wb") as f:
        writer.write(f)

    monkeypatch.setattr(sys, "argv", ["place_cv.py", "vac-789", str(src)])
    place_cv.main()

    captured = capsys.readouterr()
    assert "WARNING" in captured.out
    assert "no real selectable text" in captured.out
