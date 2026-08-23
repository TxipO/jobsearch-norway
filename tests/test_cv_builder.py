"""Tests for the deterministic CV/søknad builder. The point of the whole
rewrite (PLAN-BUILDER.md) is that structure and facts live in code, not in
model output — so these lock in the two things that kept going wrong when
the model controlled the document: self-undermining language, and the
"degree not recognized" disclaimer that's both a volunteered negative and
factually wrong."""

from docx import Document

from cv_builder import build_cv, build_cv_general, build_soknad


def _full_text(doc):
    """doc.paragraphs only covers body-level paragraphs — the photo header
    (2026-07-20) puts name/role/contact inside a table, whose cell
    paragraphs are invisible to that collection. Appending table text
    afterwards doesn't preserve reading order, but nothing here asserts
    ordering across the header/body boundary."""
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(p.text for p in cell.paragraphs)
    return "\n".join(texts)


def _cv_text(tmp_path, tailoring):
    build_cv(tailoring, tmp_path / "cv.docx")
    return _full_text(Document(str(tmp_path / "cv.docx")))


def _general_cv_text(tmp_path, tailoring):
    build_cv_general(tailoring, tmp_path / "cv-general.docx")
    return _full_text(Document(str(tmp_path / "cv-general.docx")))


def _cv_text_no(tmp_path, tailoring):
    build_cv(tailoring, tmp_path / "cv-no.docx", lang="no")
    return _full_text(Document(str(tmp_path / "cv-no.docx")))


def _general_cv_text_no(tmp_path, tailoring):
    build_cv_general(tailoring, tmp_path / "cv-general-no.docx", lang="no")
    return _full_text(Document(str(tmp_path / "cv-general-no.docx")))


def test_cv_never_says_not_recognized(tmp_path):
    text = _cv_text(tmp_path, {}).lower()
    for banned in ("not recognized", "not recognised", "not yet recogn", "ikke godkjent"):
        assert banned not in text


def test_cv_states_norwegian_degree_equivalents(tmp_path):
    text = _cv_text(tmp_path, {}).lower()
    assert "bachelorgrad" in text
    assert "fagskole" in text


def test_cv_relevant_and_other_split(tmp_path):
    text = _cv_text(tmp_path, {
        "relevant_job_ids": ["pumb", "verna"],
        "other_job_ids": ["miniso"],
    })
    assert "RELEVANT EXPERIENCE" in text
    assert "OTHER EXPERIENCE" in text
    # PUMB appears above Miniso (relevant section comes first)
    assert text.index("PUMB") < text.index("Miniso")


def test_cv_retail_tailoring_can_promote_miniso(tmp_path):
    """For a retail vacancy the tailoring flips Miniso up into relevant —
    proving the split is data-driven, not hardcoded to IT framing."""
    text = _cv_text(tmp_path, {
        "relevant_job_ids": ["miniso", "callcenter"],
        "other_job_ids": ["pumb"],
    })
    assert text.index("Miniso") < text.index("PUMB")


def test_cv_uses_default_summary_when_none_given(tmp_path):
    import profile_data as pd
    text = _cv_text(tmp_path, {})
    assert pd.DEFAULT_SUMMARY[:40] in text


def test_soknad_has_no_self_undermining_language(tmp_path):
    """Even if a tailoring step tried to smuggle in an apology, the søknad
    builder only ever emits the paragraphs it's given — but we assert on a
    clean tailoring to document the expected shape and catch a future
    builder change that might inject boilerplate."""
    build_soknad({
        "soknad": {
            "position_line": "Application for IT-konsulent, Test AS",
            "paragraphs": [
                "I am applying for the IT-konsulent role. My background at PUMB covers internal support across multiple sites.",
                "I would welcome a conversation about the role.",
            ],
        }
    }, tmp_path / "soknad.docx")
    doc = Document(str(tmp_path / "soknad.docx"))
    text = "\n".join(p.text for p in doc.paragraphs).lower()
    for banned in ("i recognize", "falls short", "i should note", "upfront", "do not yet", "a1"):
        assert banned not in text
    # structural guarantees the builder adds itself
    assert "dear hiring team," in text
    assert "kind regards," in text


def test_norwegian_soknad_uses_med_vennlig_hilsen_no_dear_greeting(tmp_path):
    """2026-07-20, user-requested NO søknad variant. Per cv-reference.md's
    own research (section 4), Norwegian formal letters don't have a "Dear
    ___,"-style opener — the letter goes straight from the position-line
    heading into the body — and close with "Med vennlig hilsen,", not a
    translated "Kind regards,"."""
    build_soknad({
        "soknad": {
            "position_line": "Søknad på stilling som IT-konsulent, Test AS",
            "paragraphs": ["Jeg søker på stillingen som IT-konsulent."],
        }
    }, tmp_path / "soknad-no.docx", lang="no")
    doc = Document(str(tmp_path / "soknad-no.docx"))
    text = "\n".join(p.text for p in doc.paragraphs).lower()
    assert "med vennlig hilsen," in text
    assert "dear" not in text
    assert "kind regards" not in text


def test_soknad_shares_cv_margins_and_header_layout(tmp_path):
    """User-requested 2026-07-19: søknad used to have its own one-off
    margins (0.7in/1.1in) while the CV used 0.5in/0.9in — a paired CV and
    søknad read back to back by the same employer should visibly belong to
    the same document set. Locks in the shared margins, and that the name
    is still the very first, bold, larger-than-body run (same visual
    hierarchy as the CV's header)."""
    build_soknad({
        "soknad": {"position_line": "Application for IT-konsulent, Test AS", "paragraphs": ["Body text here."]}
    }, tmp_path / "soknad.docx")
    doc = Document(str(tmp_path / "soknad.docx"))

    section = doc.sections[0]
    from docx.shared import Inches
    assert section.top_margin == Inches(0.5)
    assert section.left_margin == Inches(0.9)

    first_run = doc.paragraphs[0].runs[0]
    assert first_run.bold
    assert first_run.font.size.pt == 18


def test_cv_shows_hobbies_from_personal_json(tmp_path):
    """personal.json now has real hobbies (2026-07-20, synced from the
    user's own updated reference resumes) — this locks in that non-empty
    hobbies actually render as an Interests section, replacing the old
    empty-hobbies assertion that stopped matching real data."""
    text = _cv_text(tmp_path, {})
    assert "INTERESTS" in text
    assert "Martial arts (MMA)" in text


def test_norwegian_cv_shows_hobbies_no_not_english(tmp_path):
    """hobbies_no is a separate field from hobbies (2026-07-20) — the NO CV
    must use the Norwegian phrasing, not the English list untranslated."""
    text = _cv_text_no(tmp_path, {})
    assert "Kampsport (MMA)" in text
    assert "Martial arts (MMA)" not in text


def test_cv_header_uses_photo_table_when_photo_path_exists(tmp_path):
    """personal.json's photo_path now points at a real file (2026-07-20) —
    the header should switch to the 3-column photo/name/contact table
    layout instead of the old plain two-paragraph header."""
    build_cv({}, tmp_path / "cv.docx")
    doc = Document(str(tmp_path / "cv.docx"))
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.rows[0].cells) == 3


def test_cv_header_falls_back_to_plain_when_no_photo(tmp_path, monkeypatch):
    """Without a resolvable photo_path, the header must render as plain
    paragraphs (no table) — covers the branch a real personal.json with a
    photo no longer exercises."""
    import json
    import cv_builder

    fake_personal = tmp_path / "personal.json"
    fake_personal.write_text(json.dumps({
        "name": "Test User", "phone": "", "email": "", "address_line": "",
        "linkedin": "", "photo_path": "", "hobbies": [],
    }), encoding="utf-8")
    monkeypatch.setattr(cv_builder, "PERSONAL_PATH", fake_personal)

    build_cv({}, tmp_path / "cv.docx")
    doc = Document(str(tmp_path / "cv.docx"))
    assert len(doc.tables) == 0
    assert doc.paragraphs[0].runs[0].text == "Test User"


def test_cv_role_headline_differs_by_variant(tmp_path):
    """The IT-focused and general CVs show a different role headline under
    the name (2026-07-20, matching the user's own reference resumes) —
    locks in that build_cv and build_cv_general don't share one line."""
    it_text = _cv_text(tmp_path, {})
    general_text = _general_cv_text(tmp_path, {})
    assert "IT SUPPORT SPECIALIST" in it_text
    assert "CUSTOMER SERVICE & TECHNICAL SUPPORT" in general_text


def test_default_relevant_experience_is_reverse_chronological(tmp_path):
    """User-requested fix, 2026-07-17: within each section, most recent job
    first — standard CV convention. Self-employed repair (May 2022-Oct
    2024) is more recent than PUMB (Nov 2021-May 2022), which is more
    recent than Verna (May-Nov 2021, confirmed 2026-07-20)."""
    text = _cv_text(tmp_path, {})
    idx_repair = text.index("Computer & Laptop Repair Technician")
    idx_pumb = text.index("Technical Support Specialist")
    idx_verna = text.index("Field Technician")
    assert idx_repair < idx_pumb < idx_verna


def test_tools_include_ai_automation(tmp_path):
    text = _cv_text(tmp_path, {})
    assert "AI-assisted process automation" in text


def test_cv_has_hard_and_soft_skills_sections(tmp_path):
    """User-requested split, 2026-07-17: Tools became Hard Skills / Soft
    Skills — soft skills were entirely missing before. Each soft skill is
    grounded in a real fact from profile.md, same honesty rule as the rest
    of the CV (see profile_data.SOFT_SKILLS comments)."""
    import profile_data as pd
    text = _cv_text(tmp_path, {})
    assert "HARD SKILLS" in text
    assert "SOFT SKILLS" in text
    assert "TOOLS" not in text
    for skill in pd.SOFT_SKILLS:
        assert skill in text


def test_general_cv_has_no_relevant_other_split(tmp_path):
    """The broad-applications variant (retail/warehouse) — single
    chronological Experience section, no Relevant/Other framing, since
    splitting a non-technical application's timeline would look like
    padding rather than focus."""
    text = _general_cv_text(tmp_path, {})
    assert "RELEVANT EXPERIENCE" not in text
    assert "OTHER EXPERIENCE" not in text
    assert "EXPERIENCE" in text


def test_general_cv_lists_all_jobs_reverse_chronological(tmp_path):
    text = _general_cv_text(tmp_path, {})
    idx_crypto = text.index("Financial Literacy & Blockchain Instructor")
    idx_repair = text.index("Computer & Laptop Repair Technician")
    idx_pumb = text.index("Technical Support Specialist")
    idx_miniso = text.index("Sales Assistant")
    idx_callcenter = text.index("Call Center Operator")
    assert idx_crypto < idx_repair < idx_pumb < idx_miniso < idx_callcenter


def test_general_cv_uses_general_summary_by_default(tmp_path):
    import profile_data as pd
    text = _general_cv_text(tmp_path, {})
    assert pd.GENERAL_SUMMARY[:40] in text
    assert "IT support professional" not in text


def test_job_title_lines_fit_before_the_date_tab_stop():
    """Live bug caught in the rendered PDF, 2026-07-17: a long
    'Title, Company' string (73 chars — PUMB's entry with a "third-largest
    bank in Ukraine" aside jammed into the title) pushed past the
    right-aligned date tab stop, so the date collapsed with zero space
    against the title instead of aligning right. 71 chars was already
    borderline (visibly tight in the render). Threshold set well under
    both observed points, not just under the exact failure — font-width
    varies by character, so character count is a heuristic, not an exact
    measurement, and needs margin on both sides."""
    import profile_data as pd

    for jid, job in pd.JOBS.items():
        combined = f"{job['title']}, {job['company']}"
        assert len(combined) <= 65, (
            f"{jid}: '{combined}' is {len(combined)} chars — move descriptive "
            f"context (e.g. employer aside) to the location line instead"
        )
        # Norwegian compounds words, so a translated title can independently
        # blow the same budget even when the English one fits — caught live
        # with crypto_teaching's title_no on first draft (71 chars).
        combined_no = f"{job['title_no']}, {job['company']}"
        assert len(combined_no) <= 65, (
            f"{jid}: '{combined_no}' is {len(combined_no)} chars (Norwegian) — "
            f"shorten title_no or move context to the location line"
        )


def test_norwegian_cv_uses_norwegian_section_headings(tmp_path):
    """User-requested, 2026-07-17: cv-reference.md itself recommends having
    both an English and a Norwegian CV — this locks in that lang="no"
    actually renders Norwegian, not just a copy of the English structure."""
    text = _cv_text_no(tmp_path, {})
    assert "RELEVANT ERFARING" in text
    assert "ANNEN ERFARING" in text
    assert "UTDANNING" in text
    assert "FAGLIGE FERDIGHETER" in text
    assert "PERSONLIGE EGENSKAPER" in text
    assert "SPRÅK" in text
    # No English headings leaking through
    assert "RELEVANT EXPERIENCE" not in text
    assert "EDUCATION" not in text


def test_norwegian_cv_uses_norwegian_facts_not_english(tmp_path):
    import profile_data as pd
    text = _cv_text_no(tmp_path, {})
    assert pd.DEFAULT_SUMMARY_NO[:40] in text
    assert pd.DEFAULT_SUMMARY[:40] not in text
    for skill in pd.SOFT_SKILLS_NO:
        assert skill in text
    for skill in pd.SOFT_SKILLS:
        assert skill not in text
    assert "Morsmål" in text
    assert "Native" not in text


def test_norwegian_cv_localizes_present_to_naa(tmp_path):
    """Live bug found 2026-07-19 (a user asking for a compact PDF variant
    surfaced it): job['dates'] is treated as language-agnostic (numbers/
    month abbreviations don't need translation), but "present" is an
    actual English word — the NO-language CV was showing "Nov 2024 –
    present" verbatim, in both this docx pipeline and a one-off HTML
    export. Confirmed via `python -c` against the real master-cv-no.docx
    before this fix."""
    text = _cv_text_no(tmp_path, {})
    assert "present" not in text.lower()
    assert "Nov 2024 – nå" in text


def test_norwegian_general_cv_uses_norwegian_experience_heading(tmp_path):
    text = _general_cv_text_no(tmp_path, {})
    assert "ERFARING" in text
    assert "RELEVANT ERFARING" not in text
    assert "ANNEN ERFARING" not in text


def test_norwegian_cv_still_states_norwegian_degree_equivalents(tmp_path):
    """The Norwegian-language CV needs the same honest degree-equivalence
    framing as the English one — no "not recognized" language, ever."""
    text = _cv_text_no(tmp_path, {}).lower()
    assert "bachelorgrad" in text
    for banned in ("not recognized", "not recognised", "ikke godkjent", "ikke anerkjent"):
        assert banned not in text
