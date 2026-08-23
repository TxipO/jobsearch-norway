"""Files a CV the user hand-crafted/exported elsewhere (e.g. Claude Design)
under a specific vacancy, so generate_documents.py's søknad step can source
its facts from it instead of the generic profile.md (see resume_prompt.py).

User-requested 2026-07-20: the user will keep redesigning/tailoring CVs
per vacancy rather than relying on one rarely-changing master CV. Vacancy
identification is always given by the user in the same message (their own
choice over filename-encoding or defaulting to the master CV — simplest,
no guessing) — this script takes an already-resolved uuid, it does no
fuzzy matching itself.

Usage:
    py place_cv.py <uuid> <file_path>

Copies <file_path> to profile/generated/<uuid>/cv.<ext> as-is (no format
conversion — a dropped PDF stays PDF, a dropped docx stays docx), and
extracts its text into cv.txt, which becomes the fact source for that
vacancy's søknad (see resume_prompt.py's per-vacancy cv.txt lookup).
"""

import shutil
import sys
from pathlib import Path

import db

OUT_ROOT = Path(__file__).parent / "profile" / "generated"

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

# Matches resume_prompt.py's own threshold for treating cv.txt as real CV
# facts vs. a failed extraction. Below this, place_cv used to write an
# almost-empty cv.txt with no warning — the user only found out later,
# indirectly, when the resulting søknad had nothing real to draw from
# (live bug caught 2026-07-20: a PDF with no real selectable text layer,
# e.g. vector-outlined fonts, extracted to 2 bytes via pypdf).
MIN_EXTRACTED_TEXT_LENGTH = 200


def extract_text(path: Path) -> str:
    if path.suffix.lower() == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if path.suffix.lower() == ".docx":
        from docx import Document
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.extend(p.text for p in cell.paragraphs)
        return "\n".join(parts)
    raise ValueError(f"unsupported CV file type: {path.suffix} (expected .pdf or .docx)")


def place_cv(uuid: str, source_path: Path) -> Path:
    if source_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"unsupported CV file type: {source_path.suffix} (expected .pdf or .docx)")

    conn = db.connect()
    vacancy = db.get_vacancy(conn, uuid)
    if vacancy is None:
        raise ValueError(f"no vacancy with uuid {uuid!r} in the database — resolve it before calling place_cv")

    # Same slug sanitization generate_documents.py and resume_prompt.py use
    # for the folder name — using the raw uuid here instead would silently
    # place the CV somewhere generate_documents.py never looks if a uuid
    # ever contains characters the slug strips (latent inconsistency fixed
    # 2026-07-20, found while auditing this script after the encoding bug).
    slug = "".join(c if c.isalnum() or c == "-" else "-" for c in uuid.lower()).strip("-") or "job"
    out_dir = (OUT_ROOT / slug).resolve()
    if not out_dir.is_relative_to(OUT_ROOT.resolve()):
        # uuid comes from the vacancies table, which is sourced from
        # external feeds (NAV/Jobbnorge/finn) — same defense-in-depth check
        # web/app.py's download_document() already applies before touching
        # the filesystem with an externally-sourced path component.
        raise ValueError(f"resolved path for uuid {uuid!r} escapes {OUT_ROOT} — refusing")
    out_dir.mkdir(parents=True, exist_ok=True)

    dest = out_dir / f"cv{source_path.suffix.lower()}"
    shutil.copyfile(source_path, dest)

    text = extract_text(dest)
    (out_dir / "cv.txt").write_text(text, encoding="utf-8")

    return dest


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    if len(sys.argv) != 3:
        sys.exit("usage: py place_cv.py <uuid> <file_path>")
    uuid, file_path = sys.argv[1], Path(sys.argv[2])
    if not file_path.exists():
        sys.exit(f"file not found: {file_path}")

    try:
        dest = place_cv(uuid, file_path)
    except ValueError as e:
        sys.exit(str(e))

    print(f"CV placed:  {dest}")
    cv_txt = dest.parent / "cv.txt"
    extracted_length = len(cv_txt.read_text(encoding="utf-8").strip())
    if extracted_length < MIN_EXTRACTED_TEXT_LENGTH:
        print(
            f"WARNING: only {extracted_length} characters of text extracted from {dest.name} — "
            "this PDF/docx likely has no real selectable text layer (common for vector-outlined "
            "fonts in some export pipelines). The søknad step will fall back to profile.md instead "
            "of using this CV's actual content. If you have a .docx version of this CV, place that "
            "instead — python-docx extraction is far more reliable than pypdf's."
        )
    else:
        print(f"Text extracted: {cv_txt}  ({extracted_length} chars)")


if __name__ == "__main__":
    main()
